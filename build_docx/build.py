# -*- coding: utf-8 -*-
"""
Build BAB IV & BAB V (skripsi SIPADES) as .docx matching PROPOSAL-SKRIPSI.docx
formatting exactly: A4, margins 4/3/3/3 cm, 1.5 line spacing (line=360),
TNR, Heading1 (BAB) centered bold, Heading2/3 numbered, justified body with
firstLine indent, centered figure captions, shaded table headers, PAGE footer.
"""
import os, json, glob
from PIL import Image
from docx import Document
from docx.shared import Twips, Pt, Inches, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
CONTENT_DIR = os.path.join(BASE, "content")
IMG_DIR = os.path.abspath(os.path.join(BASE, "..", "screenshots"))
OUT = os.path.abspath(os.path.join(BASE, "..", "screenshots", "BAB-IV-V-SIPADES.docx"))

TBL_COUNTER = {"n": 0}

def next_table_num():
    TBL_COUNTER["n"] += 1
    return "4." + str(TBL_COUNTER["n"])

# ---------------- low-level formatting helpers ----------------

def _child(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def set_spacing(p, before=None, after=0, line=360, rule="auto"):
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), rule)

def set_indent(p, firstLine=None, left=None, hanging=None):
    pPr = p._p.get_or_add_pPr()
    ind = _child(pPr, "w:ind")
    if firstLine is not None:
        ind.set(qn("w:firstLine"), str(firstLine))
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if hanging is not None:
        ind.set(qn("w:hanging"), str(hanging))

def style_run(r, pt=12, bold=False, color=None, name="Times New Roman"):
    r.font.name = name
    rpr = r._r.get_or_add_rPr()
    rf = _child(rpr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)
    r.font.size = Pt(pt)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)

# ---------------- page setup ----------------

def setup_page(doc):
    """A4, margins: top/right/bottom=3cm, left=4cm; footer distance 1.27cm."""
    sec = doc.sections[0]
    sec.page_width  = Twips(11907)
    sec.page_height = Twips(16840)
    sec.top_margin    = Twips(1701)
    sec.right_margin  = Twips(1701)
    sec.bottom_margin = Twips(1701)
    sec.left_margin   = Twips(2268)
    sec.footer_distance = Twips(720)
    sec.header_distance = Twips(720)

# ---------------- footer (centered PAGE field) ----------------

def add_page_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # SDT wrapper matching footer1.xml
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    sdtId = OxmlElement("w:id"); sdtId.set(qn("w:val"), "-264543138")
    sdtPr.append(sdtId)
    dpo = OxmlElement("w:docPartObj")
    dpg = OxmlElement("w:docPartGallery"); dpg.set(qn("w:val"), "Page Numbers (Bottom of Page)")
    dpu = OxmlElement("w:docPartUnique")
    dpo.append(dpg); dpo.append(dpu); sdtPr.append(dpo)
    sdt.append(sdtPr)
    sdtEnd = OxmlElement("w:sdtEndPr")
    rpr_end = OxmlElement("w:rPr"); np_end = OxmlElement("w:noProof"); rpr_end.append(np_end)
    sdtEnd.append(rpr_end); sdt.append(sdtEnd)
    sdtContent = OxmlElement("w:sdtContent")

    p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "Footer"); pPr.append(pStyle)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    p_el.append(pPr)
    def fld_run(ftype):
        r = OxmlElement("w:r"); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), ftype); r.append(fc); return r
    p_el.append(fld_run("begin"))
    r_instr = OxmlElement("w:r"); it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = " PAGE   \\* MERGEFORMAT "; r_instr.append(it); p_el.append(r_instr)
    p_el.append(fld_run("separate"))
    r_val = OxmlElement("w:r"); rpr_v = OxmlElement("w:rPr"); np_v = OxmlElement("w:noProof")
    rpr_v.append(np_v); r_val.append(rpr_v)
    t_v = OxmlElement("w:t"); t_v.text = "1"; r_val.append(t_v); p_el.append(r_val)
    p_el.append(fld_run("end"))
    sdtContent.append(p_el); sdt.append(sdtContent)
    footer._element.append(sdt)

# ---------------- paragraph factories ----------------

def add_bab_heading(doc, roman, title):
    """Heading1: 'BAB IV\nHASIL DAN PEMBAHASAN' centered, TNR 12pt bold."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r1 = p.add_run("BAB " + roman); style_run(r1, pt=12, bold=True)
    r_br = p.add_run(); r_br._r.append(OxmlElement("w:br"))
    r2 = p.add_run(title); style_run(r2, pt=12, bold=True)
    return p

def add_h2(doc, num, title):
    """Heading2: '4.1  Hasil' — TNR 12pt bold, before=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    pPr = p._p.get_or_add_pPr()
    _child(pPr, "w:numPr")  # ensure exists but val=0 (no list)
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "360"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_h3(doc, num, title):
    """Heading3: '4.1.1  Lingkungan Implementasi' — TNR 12pt bold, ind left=990 hanging=720."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "240"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    ind = _child(pPr, "w:ind")
    ind.set(qn("w:left"), "990"); ind.set(qn("w:hanging"), "720")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_body(doc, text):
    """Normal body: justified, firstLine=270, TNR 12pt, line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_spacing(p, before=None, after=0, line=360)
    set_indent(p, firstLine=270)
    p.alignment = AL.JUSTIFY
    r = p.add_run(text); style_run(r, pt=12)
    return p

def add_caption(doc, text):
    """Caption: centered, TNR 11pt, before=240 after=60 line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "240"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    ind = _child(pPr, "w:ind"); ind.set(qn("w:firstLine"), "0")
    r = p.add_run(text); style_run(r, pt=11, color="000000")
    return p

# ---------------- image insertion ----------------

IMG_COUNTER = {"n": 0}

def add_figure(doc, filename, caption_text):
    IMG_COUNTER["n"] += 1
    path = os.path.join(IMG_DIR, filename)
    try:
        with Image.open(path) as im:
            w_px, h_px = im.size
    except Exception:
        w_px, h_px = 1280, 720
    # fit within 14.39 cm (usable width: 21 - 4 - 3 = 14 cm → 14.39 cm = 5.665 in)
    max_w_emu = int(5.665 * 914400)
    max_h_emu = int(3.5 * 914400)
    ratio = w_px / h_px
    w_emu = max_w_emu
    h_emu = int(w_emu / ratio)
    if h_emu > max_h_emu:
        h_emu = max_h_emu; w_emu = int(h_emu * ratio)
    p = doc.add_paragraph()
    set_spacing(p, before=120, after=0, line=360)
    p.alignment = AL.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(w_emu), height=Emu(h_emu))
    add_caption(doc, caption_text)
    return p

# ---------------- black-box test table ----------------

def add_bb_table(doc, tbl_num, caption_text, rows):
    """rows = list of (no, skenario, input, expected, result, status)"""
    add_caption(doc, f"Tabel {tbl_num} {caption_text}")
    tbl = doc.add_table(rows=1 + len(rows), cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Twips(480), Twips(1440), Twips(1440), Twips(1800), Twips(1800), Twips(720)]
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w
    headers = ["No.", "Skenario Uji", "Data Masukan", "Hasil yang Diharapkan", "Hasil Pengujian", "Status"]
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = AL.CENTER
        r = p.add_run(h); style_run(r, pt=10, bold=True)
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = AL.CENTER if ci in (0, 5) else AL.JUSTIFY
            r = p.add_run(str(val)); style_run(r, pt=10)
    return tbl
